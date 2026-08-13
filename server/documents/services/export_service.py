import io
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os
import cloudinary.uploader
from insights.models import GeneratedContractDraft

from django.conf import settings

FONT_DIR = settings.BASE_DIR / 'fonts'
try:
    pdfmetrics.registerFont(TTFont('InstrumentSerif', str(FONT_DIR / 'InstrumentSerif-Regular.ttf')))
    pdfmetrics.registerFont(TTFont('Inter', str(FONT_DIR / 'Inter-Regular.ttf')))
    pdfmetrics.registerFont(TTFont('SpaceGrotesk', str(FONT_DIR / 'SpaceGrotesk-Regular.ttf')))
    pdfmetrics.registerFont(TTFont('SpaceGrotesk-Bold', str(FONT_DIR / 'SpaceGrotesk-Bold.ttf')))
    
    FONT_TITLE = 'InstrumentSerif'
    FONT_BODY = 'Inter'
    FONT_TAG = 'SpaceGrotesk'
    FONT_TAG_BOLD = 'SpaceGrotesk-Bold'
except Exception as e:
    print(f"Error loading fonts: {e}")
    FONT_TITLE = 'Helvetica'
    FONT_BODY = 'Helvetica'
    FONT_TAG = 'Helvetica'
    FONT_TAG_BOLD = 'Helvetica-Bold'


def upload_to_cloudinary(file_buffer, filename, resource_type="raw"):
    file_buffer.seek(0)
    response = cloudinary.uploader.upload(
        file_buffer,
        resource_type=resource_type,
        public_id=filename,
        use_filename=True,
        unique_filename=True
    )
    return response.get('public_id'), response.get('secure_url')


def generate_analysis_report_pdf(document):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    styles = getSampleStyleSheet()
    
    PRIMARY_COLOR = colors.HexColor('#1e3a8a') 
    WARNING_COLOR = colors.HexColor('#f59e0b')
    DANGER_COLOR = colors.HexColor('#ef4444')  
    SAFE_COLOR = colors.HexColor('#10b981')    
    TEXT_COLOR = colors.HexColor('#333333')
    MUTED_COLOR = colors.HexColor('#64748b')

    # Custom styles
    title_style = ParagraphStyle(
        'MainTitle',
        parent=styles['Heading1'],
        fontName=FONT_TITLE,
        fontSize=28,
        leading=34,
        textColor=PRIMARY_COLOR,
        alignment=1,
        spaceAfter=15
    )
    
    subtitle_style = ParagraphStyle(
        'Subtitle',
        parent=styles['Heading2'],
        fontName=FONT_TITLE,
        fontSize=20,
        leading=26,
        textColor=PRIMARY_COLOR,
        spaceAfter=12
    )

    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontName=FONT_BODY,
        fontSize=10,
        textColor=TEXT_COLOR,
        leading=16
    )
    
    muted_style = ParagraphStyle(
        'Muted',
        parent=normal_style,
        textColor=MUTED_COLOR,
        fontSize=9,
        alignment=1
    )
    
    elements = []
    
    elements.append(Paragraph("KupasKontrak", ParagraphStyle('Logo', fontName=FONT_TITLE, fontSize=20, textColor=PRIMARY_COLOR)))
    elements.append(Spacer(1, 10))
    elements.append(HRFlowable(width="100%", thickness=1, color=PRIMARY_COLOR, spaceBefore=0, spaceAfter=20))
    
    elements.append(Paragraph("Hasil Analisis Risiko Kontrak", title_style))
    elements.append(Paragraph(f"Dokumen: <b>{document.original_filename}</b>", muted_style))
    elements.append(Spacer(1, 30))
    
    score = round(document.overall_risk_score or 0)
    
    if score >= 70:
        risk_status = "RISIKO TINGGI"
        score_color = DANGER_COLOR
    elif score >= 40:
        risk_status = "RISIKO SEDANG"
        score_color = WARNING_COLOR
    else:
        risk_status = "RISIKO RENDAH"
        score_color = SAFE_COLOR
        
    score_style = ParagraphStyle(
        'Score',
        parent=styles['Heading1'],
        fontName=FONT_TITLE,
        fontSize=48,
        leading=56,
        textColor=score_color,
        alignment=1,
        spaceAfter=15
    )
    
    status_style = ParagraphStyle(
        'Status',
        parent=styles['Normal'],
        fontName=FONT_TAG_BOLD,
        fontSize=16,
        leading=20,
        textColor=score_color,
        alignment=1,
        spaceAfter=30
    )
    
    elements.append(Paragraph(f"{score}/100", score_style))
    elements.append(Spacer(1, 15))
    elements.append(Paragraph(risk_status, status_style))
    
    elements.append(HRFlowable(width="100%", thickness=1, color=colors.lightgrey, spaceBefore=10, spaceAfter=20))
    
    elements.append(Paragraph("Rincian Klausul Berisiko", subtitle_style))
    elements.append(Spacer(1, 10))
    
    clauses = document.clauses.filter(clause_safety_score__gte=40).order_by('-clause_safety_score')
    
    if not clauses.exists():
        elements.append(Paragraph("Tidak ditemukan klausul berisiko sedang atau tinggi. Kontrak ini tergolong sangat aman.", normal_style))
    else:
        for c in clauses:
            is_critical = c.clause_safety_score >= 70
            tag_color = DANGER_COLOR if is_critical else WARNING_COLOR
            status = "KRITIS" if is_critical else "SEDANG"
            
            clause_title = ParagraphStyle('CT', fontName=FONT_TAG_BOLD, fontSize=11, textColor=tag_color, spaceAfter=8)
            elements.append(Paragraph(f"[{status}] Kategori: {c.category}", clause_title))
            
            data = [
                [Paragraph("<b>Teks Asli:</b>", normal_style), Paragraph(c.clause_text or "-", normal_style)],
                [Paragraph("<b>Analisis AI:</b>", normal_style), Paragraph(c.plain_language_summary or "-", normal_style)]
            ]
            
            if c.legal_reference:
                data.append([Paragraph("<b>Referensi Hukum:</b>", normal_style), Paragraph(c.legal_reference, normal_style)])
                
            t = Table(data, colWidths=[90, 390])
            t.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f8fafc')),
                ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#e2e8f0')),
                ('INNERGRID', (0,0), (-1,-1), 1, colors.HexColor('#e2e8f0')),
                ('BOTTOMPADDING', (0,0), (-1,-1), 8),
                ('TOPPADDING', (0,0), (-1,-1), 8),
                ('LEFTPADDING', (0,0), (-1,-1), 8),
                ('RIGHTPADDING', (0,0), (-1,-1), 8),
            ]))
            elements.append(t)
            elements.append(Spacer(1, 20))
            
    doc.build(elements)
    
    filename = f"report_{document.id}.pdf"
    public_id, secure_url = upload_to_cloudinary(buffer, filename, resource_type="raw")
    return secure_url


def generate_fixed_contract_docx(document):
    draft = document.drafts.filter(draft_type=GeneratedContractDraft.DraftType.FULL_REWRITE).order_by('-created_at').first()
    
    if not draft or not draft.content:
        raise ValueError("Draft 'full_rewrite' tidak ditemukan. Harap lakukan Auto-Fix terlebih dahulu.")
        
    content_json = draft.content
    
    docx = DocxDocument()
    
    style = docx.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    title_text = content_json.get("title", f"Revisi Kontrak: {document.original_filename}")
    heading = docx.add_heading(title_text, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sections = content_json.get("sections", [])
    
    for section in sections:
        section_title = section.get("section_title", "")
        content = section.get("content", "")
        
        if section_title:
            h2 = docx.add_heading(section_title, level=2)
            
        p = docx.add_paragraph(content)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
    docx.add_page_break()
    disclaimer_p = docx.add_paragraph()
    disclaimer_p.add_run("DISCLAIMER: ").bold = True
    disclaimer_p.add_run(draft.disclaimer)
    disclaimer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    buffer = io.BytesIO()
    docx.save(buffer)
    
    filename = f"contract_{document.id}.docx"
    public_id, secure_url = upload_to_cloudinary(buffer, filename, resource_type="raw")
    return secure_url
